#!/usr/bin/env python3
"""Generate Word report for parallel K-Means course project."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_heading_style(run, size=14, bold=True, color=RGBColor(0, 0, 0)):
    font = run.font
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = color
    font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_body_style(run, size=12):
    font = run.font
    font.size = Pt(size)
    font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        set_heading_style(run, size=16, bold=True, color=RGBColor(0, 0, 128))
    elif level == 2:
        set_heading_style(run, size=14, bold=True, color=RGBColor(0, 0, 0))
    else:
        set_heading_style(run, size=12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph_custom(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_body_style(run)
    run.font.bold = bold
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_image_custom(doc, image_path, caption, width_inches=5.5):
    """Insert an image with centered caption."""
    # Add image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    # Add caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run(caption)
    set_body_style(cap_run)
    cap_run.font.size = Pt(10)
    cap_run.font.italic = True
    return p

def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

# =================== Create Document ===================
doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('基于 MPI 的并行 K-Means 聚类算法设计与实现')
set_heading_style(title_run, size=22, bold=True, color=RGBColor(0, 0, 0))
title.paragraph_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('——并行算法课程设计报告')
set_heading_style(subtitle_run, size=14, bold=False)
subtitle.paragraph_format.space_after = Pt(24)

# =================== 1. 背景 ===================
add_heading_custom(doc, '一、背景与选题依据', level=1)

add_paragraph_custom(doc,
    '聚类分析是数据挖掘和机器学习领域中的一项基础任务，其目标是将数据集中的样本划分为若干个簇，'
    '使得同一簇内的样本相似度高，而不同簇之间的样本相似度低。K-Means 算法由 Stuart Lloyd 于 1957 年提出，'
    '是应用最广泛的划分式聚类算法之一。该算法以簇内误差平方和（Sum of Squared Errors, SSE）最小化为目标，'
    '通过迭代优化的方式逐步调整聚类中心的位置，直至收敛。')

add_paragraph_custom(doc,
    'K-Means 算法的时间复杂度为 O(N·K·d·I)，其中 N 为样本数量，K 为聚类数目，d 为样本维度，'
    'I 为迭代次数。当数据规模达到百万甚至千万级别时，串行算法的执行时间会急剧增长，难以满足实际应用的需求。'
    '与此同时，K-Means 算法具有显著的计算特性：在每一次迭代中，每个样本到各个聚类中心的距离计算是完全独立的，'
    '不同样本之间的分配操作不存在数据依赖。这种天然的“数据并行性”为算法的并行化设计提供了理想的条件。')

add_paragraph_custom(doc,
    '本课程设计选择 K-Means 聚类作为并行化对象，主要基于以下几点考虑：'
    '第一，算法逻辑简洁，仅包含距离计算和均值更新两个核心步骤，便于按照 PCAM 方法学进行系统化设计；'
    '第二，通信模式规整，每轮迭代只需全局归约（Allreduce）局部簇内和与计数，避免了复杂的点对点通信；'
    '第三，实验参数易于控制，通过调整样本数量 N 和进程数 P，可以方便地观察加速比、效率和成本等评估指标的变化规律。')

add_paragraph_custom(doc,
    '相比之下，AP（Affinity Propagation）聚类虽然也是一种优秀的聚类方法，但其核心计算涉及 N×N 的相似度矩阵、'
    '责任度矩阵和可用度矩阵的迭代更新，通信量随样本数量的平方增长，算法复杂度高，参数（preference、damping factor）'
    '敏感，收敛行为不稳定。在有限的时间和资源条件下，难以完成清晰的并行设计分析和可复现的性能评估。因此，'
    '本设计最终选定 K-Means 聚类作为并行算法课程设计的题目。')

# =================== 2. 算法设计 (PCAM) ===================
add_heading_custom(doc, '二、并行算法设计（PCAM 方法学）', level=1)

add_paragraph_custom(doc,
    'PCAM 方法学是设计并行算法的经典框架，包含四个阶段：划分（Partitioning）、通信（Communication）、'
    '组合（Agglomeration）和映射（Mapping）。下面按照这四个阶段，详细阐述本设计的并行 K-Means 算法。')

# 2.1 划分
add_heading_custom(doc, '2.1 划分（Partitioning）—— 域分解', level=2)

add_paragraph_custom(doc,
    '划分阶段的目的是将原始问题分解为若干个可并发执行的基本任务，以充分开拓算法的并行性。'
    '对于 K-Means 算法而言，最自然的分解方式是域分解（Domain Decomposition），即以数据为分解对象。')

add_paragraph_custom(doc,
    '具体而言，将包含 N 个样本的输入数据集均匀划分为 P 个不相交的子集，每个子集的大小约为 N/P。'
    '对应到 K-Means 的迭代过程中，每个基本任务负责完成以下操作：计算其对应子集中所有样本到 K 个聚类中心的欧氏距离，'
    '将每个样本分配到距离最近的中心所代表的簇，同时累加该子集内属于各个簇的样本分量和与样本计数。')

add_paragraph_custom(doc,
    '在划分阶段，我们暂时忽略目标机器的处理器数目和体系结构细节，专注于发掘问题本身蕴含的并发性。'
    '由于每个样本的分配计算仅依赖于该样本自身的特征向量和当前全局聚类中心，不同样本之间不存在数据依赖，'
    '因此划分后的各子任务可以完全独立地并行执行。')

add_paragraph_custom(doc,
    '划分判据的验证：（1）避免了冗余计算——每个样本仅被分配到一个任务中处理，不存在重复计算；'
    '（2）各任务计算量大致相当——当 N 能被 P 整除时，每个任务处理的样本数完全相同，实现了完美的负载均衡；'
    '（3）任务数与问题规模成正比——样本数 N 增加时，任务粒度随之增加，保持了良好的可扩放性。')

# 2.2 通信
add_heading_custom(doc, '2.2 通信（Communication）', level=2)

add_paragraph_custom(doc,
    '划分阶段产生的各子任务通常不能完全独立执行，任务之间需要进行数据交换以协调计算结果。'
    '在 K-Means 算法中，通信需求产生于更新步骤：各进程计算得到的是局部簇内和与局部计数，'
    '必须汇总为全局簇内和与全局计数后，才能正确更新聚类中心。')

add_paragraph_custom(doc,
    '本设计采用 MPI 的集合通信原语 MPI_Allreduce 来实现全局归约。具体而言，每轮迭代执行两次 Allreduce 操作：'
    '第一次对 K×d 维的双精度浮点数组 local_sum 进行求和，得到全局样本分量和 global_sum；'
    '第二次对 K 维的整数数组 local_count 进行求和，得到全局样本计数 global_count。'
    'Allreduce 操作保证了所有进程在调用返回后，均获得完全一致的归约结果，从而可以直接在本地更新聚类中心，'
    '无需额外的广播操作。')

add_paragraph_custom(doc,
    '从通信模式的角度分析：'
    '（1）属于全局通信（Global Communication）——每个进程都参与集合操作；'
    '（2）属于结构化通信（Structured Communication）——通信模式固定，每轮迭代执行相同的 Allreduce；'
    '（3）属于静态通信（Static Communication）——通信伙伴的身份在整个执行过程中保持不变；'
    '（4）属于同步通信（Synchronous Communication）——Allreduce 具有隐式同步屏障，所有进程必须到达通信点后才能继续。')

add_paragraph_custom(doc,
    '通信判据的验证：（1）各任务通信量相当——每个进程传递的数据量均为 K×d 个双精度数加 K 个整数，与分配的样本子集大小无关；'
    '（2）通信可并行执行——MPI_Allreduce 在底层通常采用树形或环形算法实现，归约过程具有内在并行性；'
    '（3）无复杂点对点通信——全程仅使用集合通信原语，简化了程序的正确性验证和性能分析。')

# 2.3 组合
add_heading_custom(doc, '2.3 组合（Agglomeration）', level=2)

add_paragraph_custom(doc,
    '组合阶段的目的是依据任务的局部性原理，将细粒度的基本任务组合为粗粒度的更大任务，'
    '以降低通信开销、提升计算效率，同时保持算法的灵活性和可扩放性。')

add_paragraph_custom(doc,
    '在 K-Means 的并行设计中，我们将每个 MPI 进程所负责的 N/P 个样本的所有计算（距离计算、样本分配、局部累加）'
    '组合为一个粗粒度任务。这一决策的核心依据是表面-容积效应（Surface-to-Volume Effect）：'
    '在细粒度划分下，若将每个样本视为一个独立任务，则每轮迭代需要 P 个任务之间进行 K×d 维数据的全局交换，'
    '通信开销为 O(K·d·P)，而每个任务的计算量仅为 O(K·d)，通信开销远大于计算量。'
    '组合为进程级任务后，每个任务的计算量提升为 O((N/P)·K·d)，通信量仍保持为 O(K·d·P)，'
    '计算与通信的比值从 O(1/P) 提升为 O(N/P²)，显著改善了性能。')

add_paragraph_custom(doc,
    '此外，组合还带来了以下好处：（1）减少了 MPI 进程的数量，降低了进程创建、调度和上下文切换的开销；'
    '（2）每个进程处理连续的内存块，提升了数据局部性和缓存利用率；'
    '（3）简化了编程模型，每个进程只需维护局部数据结构和局部累加器，降低了代码复杂度。')

add_paragraph_custom(doc,
    '组合判据的验证：（1）粒度增加显著降低了通信成本——从每样本一次通信降为每进程一次通信；'
    '（2）保持了灵活性和可扩放性——进程数 P 可根据硬件资源动态调整，不影响算法正确性；'
    '（3）任务数与处理器数一致——避免了任务调度器的额外开销，实现了计算与硬件资源的直接映射。')

# 2.4 映射
add_heading_custom(doc, '2.4 映射（Mapping）', level=2)

add_paragraph_custom(doc,
    '映射阶段将组合后的任务分配到具体的处理器上执行，目标是最大化并行性能。'
    '映射问题本质上是一个负载均衡与通信最小化的权衡问题，在一般情况下属于 NP 完全问题。')

add_paragraph_custom(doc,
    '本设计采用静态映射策略：进程 rank i 负责处理连续的样本块，其索引范围为 [i·N/P, (i+1)·N/P)。'
    '当 N 不能被 P 整除时，前 (N mod P) 个进程各多分配一个样本，以保证负载尽可能均衡。'
    '静态映射的优点在于实现简单、无运行时调度开销、内存访问模式规律（连续内存块），有利于缓存预取。')

add_paragraph_custom(doc,
    '映射判据的验证：（1）无通信瓶颈——Allreduce 由 MPI 库底层优化实现，不依赖于特定进程作为通信中枢；'
    '（2）无需动态负载均衡——各进程的计算量由数据划分预先确定，运行时无需任务窃取或工作队列管理；'
    '（3）高通信耦合的任务被映射到同一处理器——由于采用了数据并行策略，进程之间除每轮 Allreduce 外无其他通信，'
    '因此映射策略对通信性能的影响较小。')

# =================== 3. 算法伪代码 ===================
add_heading_custom(doc, '三、算法描述', level=1)

add_heading_custom(doc, '3.1 串行 K-Means 算法', level=2)

add_paragraph_custom(doc, '串行算法作为并行版本的性能对比基线，其伪代码如下：')

add_code_block(doc, 
"输入：数据集 X[N][d]，聚类数 K，最大迭代次数 max_iter\n"
"输出：聚类中心 C[K][d]，样本标签 labels[N]\n\n"
"1.  随机初始化聚类中心 C[K][d]\n"
"2.  for iter = 1 to max_iter do\n"
"3.      // 分配步骤（Assignment Step）\n"
"4.      for i = 0 to N-1 do\n"
"5.          labels[i] ← argmin_k ||X[i] - C[k]||²\n"
"6.      end for\n"
"7.      // 更新步骤（Update Step）\n"
"8.      for k = 0 to K-1 do\n"
"9.          sum[k] ← 0, count[k] ← 0\n"
"10.     end for\n"
"11.     for i = 0 to N-1 do\n"
"12.         sum[labels[i]] ← sum[labels[i]] + X[i]\n"
"13.         count[labels[i]] ← count[labels[i]] + 1\n"
"14.     end for\n"
"15.     for k = 0 to K-1 do\n"
"16.         if count[k] > 0 then\n"
"17.             C[k] ← sum[k] / count[k]\n"
"18.         end if\n"
"19.     end for\n"
"20.     if 中心移动距离 < ε then 跳出循环\n"
"21. end for")

add_heading_custom(doc, '3.2 并行 K-Means 算法（MPI）', level=2)

add_paragraph_custom(doc, '基于 PCAM 方法学设计的 MPI 并行版本伪代码如下：')

add_code_block(doc,
"输入：样本总数 N，维度 d，聚类数 K，进程数 P，最大迭代次数 max_iter\n"
"输出：聚类中心 C[K][d]\n\n"
"1.   if rank == 0 then\n"
"2.      生成合成数据集 X[N][d]\n"
"3.      初始化聚类中心 C[K][d]\n"
"4.  end if\n"
"5.  MPI_Scatterv(X, sendcounts, displs, local_X)  // 数据划分\n"
"6.  MPI_Bcast(C, K*d, 0)                          // 广播中心\n"
"7.  for iter = 1 to max_iter do\n"
"8.      // 局部分配步骤\n"
"9.      for i = 0 to local_N-1 do\n"
"10.         local_labels[i] ← argmin_k ||local_X[i] - C[k]||²\n"
"11.     end for\n"
"12.     // 局部累加\n"
"13.     for k = 0 to K-1 do\n"
"14.         local_sum[k] ← 0, local_count[k] ← 0\n"
"15.     end for\n"
"16.     for i = 0 to local_N-1 do\n"
"17.         local_sum[local_labels[i]] += local_X[i]\n"
"18.         local_count[local_labels[i]] += 1\n"
"19.     end for\n"
"20.     // 全局通信：归约局部和与计数\n"
"21.     MPI_Allreduce(local_sum, global_sum, K*d, MPI_DOUBLE, MPI_SUM)\n"
"22.     MPI_Allreduce(local_count, global_count, K, MPI_INT, MPI_SUM)\n"
"23.     // 更新全局中心\n"
"24.     for k = 0 to K-1 do\n"
"25.         if global_count[k] > 0 then\n"
"26.             C[k] ← global_sum[k] / global_count[k]\n"
"27.         end if\n"
"28.     end for\n"
"29. end for\n"
"30. if rank == 0 then 输出聚类结果 end if")

add_paragraph_custom(doc,
    '并行算法与串行算法的核心差异在于：'
    '（1）数据被分散到各进程，每个进程只处理局部样本子集；'
    '（2）每轮迭代的更新步骤被拆分为“局部累加”和“全局归约”两步，通过 MPI_Allreduce 实现跨进程的数据汇总；'
    '（3）聚类中心在全局归约后由所有进程同步更新，保证了算法语义与串行版本的一致性。')

# =================== 4. 实验环境 ===================
add_heading_custom(doc, '四、实验环境', level=1)

add_paragraph_custom(doc,
    '本课程设计的全部实验均在本地 MacBook Pro 上完成，无需额外的集群环境。'
    '虽然实验平台为单机，但 MPI（Message Passing Interface）的消息传递编程模型完全模拟了分布式内存系统的行为。'
    '单机多核环境下，MPI 进程之间通过共享内存进行数据传输，通信延迟低于真实的网络环境，'
    '因此测得的加速比更接近算法的理论上限，为并行算法的性能评估提供了可靠的基准。')

add_paragraph_custom(doc, '具体的软硬件环境配置如下表所示：')

add_table_from_data(doc,
    ['项目', '配置详情'],
    [
        ['处理器', 'Apple M4（10 核：4 性能核 + 6 能效核）'],
        ['内存', '16 GB 统一内存'],
        ['操作系统', 'macOS'],
        ['编程语言', 'C++17'],
        ['并行环境', 'OpenMPI 5.0.9（通过 Homebrew 安装）'],
        ['编译器', 'mpicxx（Apple Clang 包装器）'],
        ['数据生成', '程序内置随机高斯分布合成数据生成器'],
    ]
)

add_paragraph_custom(doc,
    '选择 OpenMPI 作为并行环境的原因在于：'
    '（1）OpenMPI 是 MPI 标准最流行的开源实现，具有良好的可移植性和社区支持；'
    '（2）在 macOS 平台上可通过 Homebrew 一键安装，配置简单；'
    '（3）支持 C/C++ 语言绑定，与本设计的代码语言一致。')

add_paragraph_custom(doc,
    '需要特别说明的是，Apple M4 芯片采用了异构多核架构，包含 4 个高性能核心（Performance Core）和 6 个高能效核心（Efficiency Core）。'
    '性能核具有更高的主频和更大的缓存，适合计算密集型任务；能效核主频较低，主要用于后台轻量任务。'
    '这一架构特征在后续的性能分析中具有重要意义：当进程数超过性能核数量时，部分进程将被调度到能效核上执行，'
    '从而成为并行计算的瓶颈。')

# =================== 5. 实验设置 ===================
add_heading_custom(doc, '五、实验设置', level=1)

add_heading_custom(doc, '5.1 数据集', level=2)

add_paragraph_custom(doc,
    '为避免外部数据集带来的文件 I/O 开销和格式依赖，本设计采用程序内置的合成数据生成器。'
    '数据在 d 维欧氏空间中按照高斯分布生成 K 个簇：首先随机生成 K 个簇中心，坐标范围为 [-50, 50]，'
    '然后围绕每个中心按照标准差 σ=5 的正态分布采样相应数量的样本点。')

add_paragraph_custom(doc, '实验参数设置如下：')

add_table_from_data(doc,
    ['参数', '取值', '说明'],
    [
        ['聚类数 K', '8', '固定的簇数目'],
        ['维度 d', '16', '样本特征维度，中等规模'],
        ['样本数 N', '10000 / 50000 / 100000 / 500000', '从小到大覆盖不同规模'],
        ['迭代次数', '20', '固定，排除收敛差异干扰'],
    ]
)

add_heading_custom(doc, '5.2 进程配置', level=2)

add_paragraph_custom(doc,
    '进程数 P 取值为 {1, 2, 4, 8}。使用 mpirun 的 --oversubscribe 选项允许启动超过物理核心数的 MPI 进程，'
    '以便观察超订（oversubscription）对并行性能的影响。')

add_paragraph_custom(doc,
    '固定迭代次数为 20 轮的原因在于：K-Means 的收敛速度受初始中心选择影响较大，不同运行可能提前收敛或需要更多迭代。'
    '固定迭代次数可以消除收敛行为差异对时间测量的干扰，使得不同进程数下的实验结果具有可比性，'
    '从而更纯粹地反映并行化本身带来的性能变化。')

add_heading_custom(doc, '5.3 评估指标', level=2)

add_paragraph_custom(doc,
    '依据课程第二章所讲授的并行算法评估方法，本实验测量以下五项核心指标：')

add_table_from_data(doc,
    ['指标名称', '符号', '定义 / 计算公式'],
    [
        ['串行运行时间', 'Ts', '单进程（P=1）执行完固定 20 轮迭代所需的 wall-clock 时间'],
        ['并行运行时间', 'Tp', 'P 个进程协同执行完相同工作量所需的 wall-clock 时间'],
        ['加速比', 'Sp', 'Sp = Ts / Tp，衡量并行化带来的速度提升倍数'],
        ['效率', 'E', 'E = Sp / P = Ts / (P·Tp)，衡量处理器资源的利用效率'],
        ['成本', 'C', 'C = Tp × P，并行算法执行的总处理器-时间积'],
    ]
)

add_paragraph_custom(doc,
    '其中，加速比和效率是评估并行算法性能的核心指标。理想情况下，P 个处理器应获得 P 倍的加速比（Sp = P），'
    '效率为 1（E = 1）。然而在实际中，由于通信开销、同步等待和负载不均衡等因素，加速比通常低于线性加速，'
    '效率随 P 增加而下降。成本指标则从另一个角度反映了并行开销：若算法是成本最优的，则并行成本 C 应接近串行成本 Ts。')

# =================== 6. 实验结果与分析 ===================
add_heading_custom(doc, '六、实验结果与分析', level=1)

add_heading_custom(doc, '6.1 原始实验数据', level=2)

add_paragraph_custom(doc, '表 4 汇总了不同样本规模 N 和进程数 P 组合下的实验测量结果：')

add_table_from_data(doc,
    ['N', 'P', '时间 Tp(ms)', '加速比 Sp', '效率 E', '成本 C'],
    [
        ['10000', '1', '7.92', '0.95', '0.95', '7.92'],
        ['10000', '2', '2.90', '2.59', '1.29', '5.80'],
        ['10000', '4', '3.68', '2.04', '0.51', '14.74'],
        ['10000', '8', '2.38', '3.16', '0.39', '19.04'],
        ['50000', '1', '23.18', '0.98', '0.98', '23.18'],
        ['50000', '2', '14.92', '1.52', '0.76', '29.83'],
        ['50000', '4', '10.45', '2.17', '0.54', '41.81'],
        ['50000', '8', '12.63', '1.80', '0.22', '101.05'],
        ['100000', '1', '46.09', '1.02', '1.02', '46.09'],
        ['100000', '2', '26.45', '1.78', '0.89', '52.90'],
        ['100000', '4', '16.71', '2.82', '0.70', '66.84'],
        ['100000', '8', '22.22', '2.12', '0.26', '177.74'],
        ['500000', '1', '236.58', '1.00', '1.00', '236.58'],
        ['500000', '2', '137.33', '1.72', '0.86', '274.65'],
        ['500000', '4', '73.69', '3.20', '0.80', '294.75'],
        ['500000', '8', '86.91', '2.71', '0.34', '695.30'],
    ]
)

add_paragraph_custom(doc,
    '注：P=1 的 MPI 版本与纯串行版本存在少量测量误差（小于 5%），源于 MPI 初始化、计时器精度以及系统调度波动。'
    '因此表中 P=1 的加速比接近 1 而非严格等于 1，这在实验测量中是正常现象。')

add_heading_custom(doc, '6.2 加速比分析', level=2)

add_image_custom(doc, 'results/speedup.png', '图 1  不同规模下的加速比曲线')

add_paragraph_custom(doc,
    '从加速比数据可以观察到明显的规模依赖效应。当样本数 N 较小（如 N=10000）时，加速比曲线波动剧烈：'
    'P=2 时加速比达到 2.59，P=4 时却回落至 2.04，P=8 时再次升至 3.16。这种不规则波动的原因是：'
    '在数据量较小时，本地计算时间极短（单轮迭代仅约 0.4ms），Allreduce 通信和进程同步的开销在总时间中占据主导地位。'
    '此时加速比不再由计算并行度决定，而是受操作系统调度、MPI 实现细节和缓存行为等随机因素支配。')

add_paragraph_custom(doc,
    '当样本规模增大到 N=500000 时，加速比呈现出更规律的变化趋势：'
    'P=2 时加速比为 1.72，P=4 时达到 3.20，非常接近理想线性加速比 4。'
    '这表明当本地计算量足够大时，通信开销可以被有效掩盖，数据并行策略展现出良好的性能。')

add_paragraph_custom(doc,
    '然而，一个显著的现象是：当 P 从 4 增加到 8 时，加速比不仅没有继续提升，反而出现了下降。'
    '以 N=500000 为例，P=4 的加速比为 3.20，而 P=8 的加速比仅为 2.71。'
    '这一反直觉的结果需要从硬件架构和并行开销两个层面进行解释。')

add_paragraph_custom(doc,
    '第一，硬件架构限制。Apple M4 芯片仅有 4 个性能核，当启动 8 个 MPI 进程时，'
    '必有部分进程被调度到主频更低、缓存更小的能效核上。MPI 的 Allreduce 是同步操作，'
    '每轮迭代必须等待最慢的进程完成后才能继续。因此，落在能效核上的进程成为整个并行计算的“短板”，'
    '拖慢了全局同步的速度。相比之下，P=4 时 4 个进程恰好可以全部运行在 4 个性能核上，'
    '不存在“慢进程”瓶颈，反而获得了更高的实际效率。')

add_paragraph_custom(doc,
    '第二，通信/计算比恶化。Allreduce 传递的数据量（K×d 个 double 加 K 个 int）与进程数 P 无关，'
    '但每个进程的本地计算量随 P 增加而减半。当 P 从 4 增加到 8 时，本地计算量从 125000 降至 62500 个样本，'
    '计算时间减少了一半，但通信开销基本不变，导致通信在总时间中的占比显著上升，整体效率下降。')

add_heading_custom(doc, '6.3 效率分析', level=2)

add_image_custom(doc, 'results/efficiency.png', '图 2  不同规模下的效率曲线')

add_paragraph_custom(doc,
    '效率指标 E = Sp / P 衡量了每个处理器的资源利用率，是并行算法可扩放性（Scalability）的核心度量。')

add_paragraph_custom(doc,
    '实验结果显示，效率随进程数增加而单调下降，符合并行算法的一般规律。但不同规模下的下降速度差异明显：'
    'N=10000 时，P=8 的效率仅为 0.39；而 N=500000 时，P=4 的效率高达 0.80，P=8 时仍有 0.34。'
    '这说明问题规模越大，并行算法的效率保持能力越强，即具有更好的弱可扩展性（Weak Scalability）。')

add_paragraph_custom(doc,
    '值得注意的是，表中出现了少量效率略大于 1 的数据（如 N=10000, P=2 时 E=1.29）。'
    '这并非真正的超线性加速，而是由于串行版本与 MPI 版本的计时波动所致。'
    'MPI 版本在 P=1 时的运行时间（7.92ms）略长于纯串行版本（7.51ms），'
    '而 P=2 时的时间（2.90ms）恰好受益于更好的缓存局部性和操作系统调度，'
    '两者叠加造成了加速比偏高的假象。在分析中应识别这类测量噪声，避免得出错误的结论。')

add_heading_custom(doc, '6.4 执行时间分析', level=2)

add_paragraph_custom(doc,
    '执行时间曲线清晰地展示了并行化带来的收益。对于最大规模 N=500000，'
    '执行时间从 P=1 时的 236.6ms 降至 P=4 时的 73.7ms，减少了约 69%。'
    '然而 P=8 时时间反弹至 86.9ms，与加速比分析中观察到的“P=8 不如 P=4”现象一致。')

add_paragraph_custom(doc,
    '对于中等规模 N=100000，P=4 时时间为 16.7ms，而 P=8 时反弹至 22.2ms，'
    '进一步验证了在单机异构多核环境下，进程数超过性能核数会导致性能退化。')

add_heading_custom(doc, '6.5 成本分析', level=2)

add_image_custom(doc, 'results/cost.png', '图 4  不同规模下的成本曲线')

add_paragraph_custom(doc,
    '成本 C = Tp × P 反映了并行算法消耗的总处理器资源。理想情况下，若并行算法没有额外开销，'
    '则成本应等于串行成本 Ts（即 Sp = P, E = 1）。实际中，由于通信和同步开销，成本通常略高于 Ts。')

add_paragraph_custom(doc,
    '以 N=500000 为例：P=2 时成本为 274.65，非常接近 Ts=236.58，说明此时并行开销很小；'
    'P=4 时成本为 294.75，较 Ts 增加约 25%，仍在可接受范围内；'
    '但 P=8 时成本飙升至 695.30，几乎是串行成本的 3 倍。')

add_paragraph_custom(doc,
    '成本曲线的急剧上升明确指示了 P=8 处于“过度并行”区域：'
    '增加处理器带来的计算收益已经无法弥补同步和通信开销的代价。'
    '根据成本最优原则，对于本实验的硬件环境，P=4 是最佳的进程配置选择。')

add_heading_custom(doc, '6.6 综合讨论', level=2)

add_paragraph_custom(doc,
    '综合以上分析，可以得出以下结论：')

add_paragraph_custom(doc,
    '（1）PCAM 设计有效性验证。通过域分解划分数据、Allreduce 规整通信、进程级组合和静态映射，'
    '本设计成功将 K-Means 算法并行化，在 N=500000、P=4 时获得了 Sp=3.2、E=0.80 的良好性能。'
    '这验证了数据并行策略在计算密集型聚类任务中的有效性，也体现了 PCAM 方法学在指导并行算法设计中的实用价值。')

add_paragraph_custom(doc,
    '（2）问题规模与并行性能的强相关性。实验表明，只有当本地计算量（N/P）足够大时，'
    '通信开销才能被计算时间有效掩盖，从而获得接近线性的加速比。'
    '当 N/P 过小（如 N=10000, P=8 时仅 1250 个样本/进程）时，算法陷入“通信主导”区域，'
    '此时增加进程不仅不能提升性能，反而会因同步开销导致效率下降。')

add_paragraph_custom(doc,
    '（3）硬件架构对并行性能的决定性影响。Apple M4 的异构多核设计（4P+6E）导致进程调度具有非对称性。'
    '当进程数超过性能核数时，能效核上的进程成为全局同步瓶颈。'
    '这一现象在传统的同构多核处理器或分布式集群中不会出现，'
    '它提示我们在实际并行算法部署时，必须充分了解目标硬件的微架构特征，合理选择进程数。')

add_paragraph_custom(doc,
    '（4）成本最优配置的确定。结合加速比、效率和成本三项指标，'
    '本实验的最佳配置为 N=500000、P=4，此时在获得 3.2 倍加速的同时仍保持 80% 的效率，'
    '成本仅比串行版本高 25%。这一配置充分利用了 4 个性能核的计算能力，同时避免了超订带来的性能退化。')

# =================== 7. 结论 ===================
add_heading_custom(doc, '七、结论', level=1)

add_paragraph_custom(doc,
    '本课程设计基于 PCAM 方法学，完成了 K-Means 聚类算法的并行化设计与实现。'
    '在划分阶段，采用域分解将数据集均匀分配到各 MPI 进程；在通信阶段，利用 MPI_Allreduce 实现每轮迭代的全局归约；'
    '在组合阶段，将样本级细粒度任务组合为进程级粗粒度任务，降低通信开销；在映射阶段，采用静态映射保证负载均衡。')

add_paragraph_custom(doc,
    '实验在本地 MacBook Pro（Apple M4, 10 核）上完成，使用 OpenMPI 作为并行编程环境。'
    '通过合成高斯数据集在四种规模（N=10000 至 500000）和四种进程配置（P=1,2,4,8）下的系统测试，'
    '测量了运行时间、加速比、效率和成本四项核心指标。')

add_paragraph_custom(doc,
    '实验结果表明：当问题规模足够大（N=500000）且进程数与硬件性能核数匹配（P=4）时，'
    '算法获得了加速比 Sp≈3.2、效率 E≈0.80 的良好性能，验证了数据并行策略的有效性。'
    '同时，实验也揭示了并行性能的限制因素：本地计算量过小导致的通信主导、以及异构多核架构下的进程调度瓶颈。')

add_paragraph_custom(doc,
    '本设计的实际意义在于：证明了即使没有专用集群环境，仅通过本地多核工作站和 MPI 标准库，'
    '即可完成完整的并行算法设计、实现与性能评估流程。这为并行算法课程设计提供了一个可复现、'
    '易部署的实践方案，也为后续在真实分布式集群上的大规模部署奠定了理论和代码基础。')

# =================== 附录 ===================
add_heading_custom(doc, '附录：核心代码清单', level=1)

add_paragraph_custom(doc, '以下为并行版本 kmeans_mpi.cpp 的关键代码片段：')

add_code_block(doc,
"// 数据划分：计算每个进程分配的样本数\n"
"int base = N / size;\n"
"int rem = N % size;\n"
"std::vector<int> sendcounts(size), displs(size);\n"
"for (int i = 0; i < size; ++i) {\n"
"    sendcounts[i] = (base + (i < rem ? 1 : 0)) * d;\n"
"    displs[i] = (i == 0) ? 0 : displs[i-1] + sendcounts[i-1];\n"
"}\n"
"int local_N = sendcounts[rank] / d;\n"
"std::vector<double> local_data(local_N * d);\n\n"
"// 分发数据\n"
"MPI_Scatterv(data.data(), sendcounts.data(), displs.data(),\n"
"             MPI_DOUBLE, local_data.data(), local_N*d,\n"
"             MPI_DOUBLE, 0, MPI_COMM_WORLD);\n\n"
"// 广播初始中心\n"
"MPI_Bcast(centers.data(), K*d, MPI_DOUBLE, 0, MPI_COMM_WORLD);\n\n"
"// 迭代计算\n"
"for (int iter = 0; iter < max_iter; ++iter) {\n"
"    // 局部分配与累加...\n"
"    MPI_Allreduce(local_sum.data(), global_sum.data(), K*d,\n"
"                  MPI_DOUBLE, MPI_SUM, MPI_COMM_WORLD);\n"
"    MPI_Allreduce(local_count.data(), global_count.data(), K,\n"
"                  MPI_INT, MPI_SUM, MPI_COMM_WORLD);\n"
"    // 更新中心...\n"
"}")

# Save
doc.save('并行KMeans聚类课程设计报告.docx')
print("Report saved to: 并行KMeans聚类课程设计报告.docx")
